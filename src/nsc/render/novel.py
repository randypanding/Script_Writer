"""小说渲染：Markdown 与 docx（T-08）。

小说来自 `chapters`（Pass6 产物）。每个段落经 `anchor_map[{paragraph_index, beat_id, line_ids}]`
携带锚点：docx 渲染时给每段嵌 `NID_<beat_id>` 书签 + 文末锚点索引表（D29①②）。
段落顺序即 `paragraphs` 顺序，与 IR 的 linear_index 严格一致（D29③）。
"""

from __future__ import annotations

from spec.ir.container import NarrativeIR, NovelChapter

from nsc.render.anchors import ANCHOR_APPENDIX_TITLE, append_anchor_index, embed_paragraph_bookmark


def _episode_title(ir: NarrativeIR, chapter: NovelChapter) -> str:
    ep = next((e for e in ir.episodes if e.id == chapter.episode_id), None)
    return f"第{ep.no}集 {ep.title}" if ep else chapter.episode_id


def to_novel_md(ir: NarrativeIR) -> str:
    """小说 → Markdown。段落前不嵌锚点（纯文本），锚点索引表附文末。"""
    blocks: list[str] = [f"# {ir.project.title}", ""]
    for chapter in ir.chapters:
        blocks.append(f"## {chapter.episode_id and _episode_title(ir, chapter)}")
        for para in chapter.paragraphs:
            blocks.append(para)
            blocks.append("")
        if chapter.anchor_map:
            blocks.append("")
            blocks.append(ANCHOR_APPENDIX_TITLE)
            blocks.append("| 段落序号 | 节点ID |")
            blocks.append("|---|---|")
            for am in chapter.anchor_map:
                blocks.append(f"| {am['paragraph_index']} | {am['beat_id']} |")
            blocks.append("")
    return "\n".join(blocks)


def to_novel_docx(ir: NarrativeIR) -> "Document":
    """小说 → docx（每段嵌书签 + 文末锚点索引表）。"""
    from docx import Document

    doc = Document()
    doc.add_heading(ir.project.title, level=0)
    bm_id = 0
    anchors: list[tuple[int, str]] = []
    for chapter in ir.chapters:
        doc.add_heading(_episode_title(ir, chapter), level=1)
        anchor_by_idx = {am["paragraph_index"]: am for am in chapter.anchor_map}
        for idx, para in enumerate(chapter.paragraphs):
            p = doc.add_paragraph(para)
            am = anchor_by_idx.get(idx)
            if am:
                embed_paragraph_bookmark(p, am["beat_id"], bm_id=bm_id)
                anchors.append((len(doc.paragraphs) - 1, am["beat_id"]))
                bm_id += 1
    append_anchor_index(doc, anchors)
    return doc