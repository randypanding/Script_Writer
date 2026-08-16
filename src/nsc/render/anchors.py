"""交付锚点（D29 三重方案）的写入与回收。

① docx 书签 `NID_<ulid>`（不可见，最可靠）
② 文末「锚点索引」表（段落序号 ↔ 节点ID，用户删表则失效）
③ 全部丢失时的模糊回退 → `src/nsc/feedback/align.py::align_paragraphs`（T-10）

这里只负责 ① ② 的写入与直接回收。`extract_docx_anchors` 是 L0 的 `test_anchor_roundtrip`
（渲染 → 解析回来 → 100% 恢复 node_id）的解析端。
"""

from __future__ import annotations

import re
from collections.abc import Iterator
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

NID_PREFIX = "NID_"
ANCHOR_APPENDIX_TITLE = "锚点索引（此表可删除）"
_NID_RE = re.compile(rf"^{NID_PREFIX}([0-9A-HJKMNP-TV-Z]{{26}})$")


def bookmark_name(node_id: str) -> str:
    """Word 书签名：`NID_<ulid>`。ULID 字符集不含歧义字符，安全。"""
    return f"{NID_PREFIX}{node_id}"


def _append_bookmark(paragraph, name: str, *, bm_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bm_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bm_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def embed_paragraph_bookmark(paragraph, node_id: str, *, bm_id: int) -> None:
    """在段落起始处插入不可见书签 `NID_<node_id>`。`bm_id` 全文档唯一。"""
    _append_bookmark(paragraph, bookmark_name(node_id), bm_id=bm_id)


def append_anchor_index(doc: Document, anchors: list[tuple[int, str]]) -> None:
    """文末追加「锚点索引」表：一列段落序号，一列节点 ID。标题提示可删除。"""
    heading = doc.add_paragraph()
    run = heading.add_run(ANCHOR_APPENDIX_TITLE)
    run.bold = True
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "段落序号"
    hdr[1].text = "节点ID"
    for idx, node_id in anchors:
        cells = table.add_row().cells
        cells[0].text = str(idx)
        cells[1].text = node_id


def _iter_bookmarks(doc: Document) -> Iterator[tuple[int, str]]:
    """按正文段落顺序 yield (段落序号, node_id)。遇到锚点附录标题即停。"""
    para_index = 0
    for p in doc.paragraphs:
        if p.text.strip() == ANCHOR_APPENDIX_TITLE:
            break
        for bm in p._p.iter(qn("w:bookmarkStart")):
            name = bm.get(qn("w:name")) or ""
            m = _NID_RE.match(name)
            if m:
                yield para_index, m.group(1)
        para_index += 1


def extract_bookmarks(doc: Document) -> dict[int, str]:
    """L1：从 docx 书签回收 node_id：段落序号 -> node_id。"""
    return dict(_iter_bookmarks(doc))


def _iter_appendix(doc: Document) -> Iterator[tuple[int, str]]:
    """L2：从文末「锚点索引」表回收。"""
    for table in doc.tables:
        for row in table.rows[1:]:
            cells = row.cells
            if len(cells) >= 2:
                try:
                    idx = int(cells[0].text.strip())
                except ValueError:
                    continue
                node_id = cells[1].text.strip()
                if _NID_RE.match(node_id):
                    yield idx, node_id


def extract_appendix(doc: Document) -> dict[int, str]:
    """L2：从「锚点索引」表回收 node_id：段落序号 -> node_id。"""
    return dict(_iter_appendix(doc))


def extract_docx_anchors(doc: Document) -> dict[int, str]:
    """合并 L1+L2。书签优先，附录表兜底。"""
    anchors = extract_bookmarks(doc)
    for idx, node_id in _iter_appendix(doc):
        anchors.setdefault(idx, node_id)
    return anchors


def extract_docx_anchors_path(path: str | Path) -> dict[int, str]:
    return extract_docx_anchors(Document(str(path)))