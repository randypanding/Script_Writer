"""DOCX 渲染 + L1 书签 / L2 附录锚点（D29）。

- 写：每个锚定段落写入 `NID_<ulid>` 书签（L1），文末追加锚点索引表（L2）。
- 读：解析 `w:bookmarkStart`，把 (node_id, text) 逐段读回，供反向对齐器使用。
"""

from __future__ import annotations

from pathlib import Path

import lxml.etree as etree
from docx import Document

from .anchors import AnchorIndex, Paragraph, beacon_for, node_id_from_beacon

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _bookmark_start(paragraph_id: int, name: str) -> etree._Element:
    el = etree.Element(f"{{{_W}}}bookmarkStart")
    el.set(f"{{{_W}}}id", str(paragraph_id))
    el.set(f"{{{_W}}}name", name)
    return el


def _bookmark_end(paragraph_id: int) -> etree._Element:
    el = etree.Element(f"{{{_W}}}bookmarkEnd")
    el.set(f"{{{_W}}}id", str(paragraph_id))
    return el


def render_docx(paragraphs: list[Paragraph], path: str | Path) -> Path:
    """写入带锚点的 docx。未锚定的段落（node_id=None）不写书签。"""
    doc = Document()
    counter = 1
    for para in paragraphs:
        p = doc.add_paragraph()
        if para.node_id:
            p._p.insert(0, _bookmark_start(counter, beacon_for(para.node_id)))
            p._p.append(_bookmark_end(counter))
            counter += 1
        p.add_run(para.text)

    # L2 附录：锚点索引表
    index = AnchorIndex(
        entries=[(i + 1, para.node_id) for i, para in enumerate(paragraphs) if para.node_id]
    )
    doc.add_heading("锚点索引", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "段落序号"
    table.rows[0].cells[1].text = "节点ID"
    for no, nid in index.entries:
        row = table.add_row()
        row.cells[0].text = str(no)
        row.cells[1].text = nid

    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def read_docx_anchors(path: str | Path) -> list[Paragraph]:
    """逐段读回 (node_id, text)。node_id 来自 L1 书签；无书签的段为 None。"""
    doc = Document(str(path))
    out: list[Paragraph] = []
    for p in doc.paragraphs:
        # 跳过附录标题（Heading 样式不是正文段）
        if p.style and p.style.name and p.style.name.lower().startswith("heading"):
            continue
        if p.text == "" and len(p._p.findall(f"{{{_W}}}bookmarkStart")) == 0:
            continue
        node_id: str | None = None
        for bm in p._p.iter(f"{{{_W}}}bookmarkStart"):
            name = bm.get(f"{{{_W}}}name", "")
            nid = node_id_from_beacon(name)
            if nid:
                node_id = nid
                break
        out.append(Paragraph(node_id=node_id, text=p.text, kind="novel_paragraph"))
    # 跳过附录表（无书签的表格段落我们不在正文读；锚点索引表是表格，非正文段）
    return out
