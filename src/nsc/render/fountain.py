"""剧本渲染为 Fountain 格式（T-08）。

Fountain 是纯文本剧本标准（fountain.io substrate）。这里从 IR 的
Scene → Beat → Line 派生剧本，`#` 注释行嵌入 Beat_kind/summary 供人读、
`==` 场景分隔。不携带锚点（docx 才带，见 render/docx.py）。
"""

from __future__ import annotations

from spec.ir.container import NarrativeIR
from spec.ir.overlays import Character, Location

_DAY = {"day": "日", "night": "夜", "dawn": "黎明", "dusk": "黄昏"}


def _loc_name(ir: NarrativeIR, loc_id: str, interior: bool) -> str:
    loc = next((x for x in ir.locations if x.id == loc_id), None)
    name = loc.name if loc else "未知地点"
    return f"{'内' if interior else '外'}景 {name}"


def _char_name(ir: NarrativeIR, char_id: str | None) -> str:
    if not char_id:
        return "叙述"
    ch = next((c for c in ir.characters if c.id == char_id), None)
    return ch.name if ch else "未知角色"


def _line_text(line) -> str:
    if line.line_type == "dialogue":
        return line.text
    if line.line_type == "action":
        return line.text
    if line.line_type == "voiceover":
        return f"（画外音）{line.text}"
    if line.line_type == "caption":
        return f"{line.text}"
    if line.line_type == "sfx":
        return f"（音效）{line.text}"
    return line.text


def to_fountain(ir: NarrativeIR) -> str:
    """整部剧本 → Fountain 文本。"""
    blocks: list[str] = [ir.project.title and f"Title: {ir.project.title}", "Credit: 集", ""]
    for ep in ir.episodes:
        blocks.append(f"== 第{ep.no}集 {ep.title}")
        for sc in (s for s in ir.scenes if s.parent_id == ep.id):
            blocks.append(f"{_loc_name(ir, sc.location_id, sc.interior)} - {_DAY.get(sc.time_of_day, '日')}")
            blocks.append(f"（本场：{sc.goal}｜冲突：{sc.conflict}｜转折：{sc.turn}）")
            for beat in (b for b in ir.beats if b.parent_id == sc.id):
                blocks.append(
                    f"#{beat.beat_kind} · {beat.summary}"
                )
                for line in (ln for ln in ir.lines if ln.parent_id == beat.id):
                    if line.line_type == "action":
                        blocks.append(_line_text(line))
                    elif line.line_type == "dialogue":
                        blocks += [_char_name(ir, line.character_id), _line_text(line)]
                    else:
                        blocks.append(_line_text(line))
            blocks.append("")
    return "\n".join(blocks)


def to_script_docx(ir: NarrativeIR) -> "Document":
    """剧本 → docx（含每行锚点书签）。"""
    from docx import Document

    doc = Document()
    doc.add_heading(ir.project.title, level=0)
    bm_id = 0
    anchors: list[tuple[int, str]] = []
    for ep in ir.episodes:
        doc.add_heading(f"第{ep.no}集 {ep.title}", level=1)
        for sc in (s for s in ir.scenes if s.parent_id == ep.id):
            doc.add_heading(f"{_loc_name(ir, sc.location_id, sc.interior)} - {_DAY.get(sc.time_of_day, '日')}", level=2)
            for beat in (b for b in ir.beats if b.parent_id == sc.id):
                for line in (ln for ln in ir.lines if ln.parent_id == beat.id):
                    p = doc.add_paragraph()
                    if line.line_type == "dialogue":
                        p.add_run(f"{_char_name(ir, line.character_id)}").bold = True
                        p.add_run(f"：{_line_text(line)}")
                    else:
                        p.add_run(_line_text(line))
                    from nsc.render.anchors import embed_paragraph_bookmark

                    embed_paragraph_bookmark(p, line.id, bm_id=bm_id)
                    anchors.append((len(doc.paragraphs) - 1, line.id))
                    bm_id += 1
    from nsc.render.anchors import append_anchor_index

    append_anchor_index(doc, anchors)
    return doc